import os
from typing import Dict, Any
from docx import Document
from src.backend.parsers.base import BaseParser

class DocxParser(BaseParser):
    """
    Parses Microsoft Word (.docx) files using python-docx.
    """

    def parse(self, file_path: str) -> str:
        """Extracts text from paragraphs and formats tables with markdown pipe formatting."""
        try:
            doc = Document(file_path)
            content_parts = []
            
            # Extract paragraphs
            for p in doc.paragraphs:
                if p.text.strip():
                    content_parts.append(p.text)
                    
            # Extract tables and format with pipe separators
            for table in doc.tables:
                table_lines = []
                for row in table.rows:
                    cells_text = [cell.text.replace("\n", " ").strip() for cell in row.cells]
                    table_lines.append(f"| {' | '.join(cells_text)} |")
                
                if table_lines and len(table.columns) > 0:
                    # Construct simple Markdown table divider line
                    header_divider = f"|{'|'.join(['---' for _ in range(len(table.columns))])}|"
                    table_lines.insert(1, header_divider)
                    content_parts.append("\n".join(table_lines))
                    
            return "\n\n".join(content_parts)
        except Exception as e:
            raise RuntimeError(f"Failed to parse Word document {file_path}: {str(e)}")

    def get_metadata(self, file_path: str) -> Dict[str, Any]:
        """Extracts core properties and metadata from the Word document."""
        try:
            doc = Document(file_path)
            props = doc.core_properties
            
            # Helper to format dates to ISO string
            def format_date(d) -> str:
                return d.isoformat() if d else ""

            return {
                "file_name": os.path.basename(file_path),
                "file_size_bytes": os.path.getsize(file_path),
                "file_type": "docx",
                "title": props.title or "",
                "author": props.author or "",
                "created": format_date(props.created),
                "modified": format_date(props.modified),
                "category": props.category or "",
                "comments": props.comments or ""
            }
        except Exception as e:
            # Fallback to basic parameters on parsing failures
            return {
                "file_name": os.path.basename(file_path),
                "file_size_bytes": os.path.getsize(file_path),
                "file_type": "docx",
                "error": f"Metadata extraction failed: {str(e)}"
            }
